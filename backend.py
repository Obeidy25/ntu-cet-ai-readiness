"""
NTU CET - RAG Backend Server
Provides PDF, Excel, and DOCX ingestion, ChromaDB vector storage, semantic sentence-aware chunking,
query routing, dynamic conversational onboarding, and hybrid multi-provider LLM generation.
"""

import os
import re
import fitz
import requests
import base64
import io
import sqlite3
import json
import asyncio
from datetime import datetime
from PIL import Image
from typing import Optional, List, Dict
from openai import OpenAI
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import tempfile
import chromadb
import openpyxl
from docx import Document as DocxDocument

load_dotenv()

app = FastAPI(title="RAG API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

chroma_client = chromadb.PersistentClient(path="./chroma_db")


def get_or_create_collection():
    """Retrieves or initializes the ChromaDB vector collection."""
    try:
        return chroma_client.get_collection("rag_collection")
    except Exception:
        return chroma_client.create_collection("rag_collection")


# === Concurrency Control & Thread-Safety Locks ===
chroma_write_lock = asyncio.Lock()
sessions_lock = asyncio.Lock()
generation_semaphore = asyncio.Semaphore(2)


# === Persistent Chat Sessions SQLite Database ===
DB_PATH = "./chat_sessions.db"


def init_sessions_db():
    """Initializes the SQLite database with WAL mode for safe multi-request concurrency."""
    conn = sqlite3.connect(DB_PATH, timeout=30.0)
    cursor = conn.cursor()
    cursor.execute("PRAGMA journal_mode=WAL;")
    cursor.execute("PRAGMA busy_timeout=10000;")
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS sessions (
            id TEXT PRIMARY KEY,
            title TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            messages TEXT
        )
    """)
    conn.commit()
    conn.close()


init_sessions_db()


def get_all_sessions() -> List[Dict]:
    """Retrieves all chat sessions sorted by most recent activity."""
    conn = sqlite3.connect(DB_PATH, timeout=30.0)
    cursor = conn.cursor()
    cursor.execute("SELECT id, title, created_at, updated_at, messages FROM sessions ORDER BY updated_at DESC")
    rows = cursor.fetchall()
    conn.close()

    sessions = []
    for row in rows:
        session_id, title, created_at, updated_at, messages_json = row
        try:
            msgs = json.loads(messages_json) if messages_json else []
        except Exception:
            msgs = []
        sessions.append({
            "id": session_id,
            "title": title or "Untitled Chat",
            "created_at": created_at,
            "updated_at": updated_at,
            "message_count": len(msgs),
        })
    return sessions


def get_session_by_id(session_id: str) -> Optional[Dict]:
    """Retrieves a specific chat session with its full message history."""
    conn = sqlite3.connect(DB_PATH, timeout=30.0)
    cursor = conn.cursor()
    cursor.execute("SELECT id, title, created_at, updated_at, messages FROM sessions WHERE id = ?", (session_id,))
    row = cursor.fetchone()
    conn.close()

    if not row:
        return None

    session_id, title, created_at, updated_at, messages_json = row
    try:
        msgs = json.loads(messages_json) if messages_json else []
    except Exception:
        msgs = []

    return {
        "id": session_id,
        "title": title,
        "created_at": created_at,
        "updated_at": updated_at,
        "messages": msgs,
    }


def save_or_update_session(session_id: str, title: str, messages: List[Dict]):
    """Saves or updates a conversation session in SQLite."""
    conn = sqlite3.connect(DB_PATH, timeout=30.0)
    cursor = conn.cursor()
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    messages_json = json.dumps(messages, ensure_ascii=False)

    cursor.execute("""
        INSERT INTO sessions (id, title, created_at, updated_at, messages)
        VALUES (?, ?, ?, ?, ?)
        ON CONFLICT(id) DO UPDATE SET
            title = excluded.title,
            updated_at = excluded.updated_at,
            messages = excluded.messages
    """, (session_id, title, now_str, now_str, messages_json))
    conn.commit()
    conn.close()


def delete_session_by_id(session_id: str) -> bool:
    """Deletes a chat session from SQLite."""
    conn = sqlite3.connect(DB_PATH, timeout=30.0)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM sessions WHERE id = ?", (session_id,))
    deleted = cursor.rowcount > 0
    conn.commit()
    conn.close()
    return deleted


CHUNK_SIZE = 800
CHUNK_OVERLAP = 100

ITU_AI_READINESS_DIMENSIONS = [
    "Data/model Marketplace",
    "Generated Content Marketplace",
    "Cross-domain Correlation Analysis",
    "Contextualization & Regional Impact",
    "Level of Integration of AI in Workflows",
    "Human Interface",
    "Strategy Alignment",
    "Collaboration with AI",
    "Impacts of Humans in AI Integration",
    "AI & Policies",
    "AI for Inclusion",
    "Granular Priorities",
    "Digital Infrastructure",
]


def table_to_markdown_chunks(table_data, max_rows=25):
    """
    Converts a 2D list of extracted cell strings into one or more Markdown tables.
    Large tables are split into multiple chunks (default max 25 rows) while preserving the header.
    Returns: List of Markdown table strings.
    """
    if not table_data or not table_data[0]:
        return []

    cleaned_rows = []
    for row in table_data:
        cleaned_row = [" ".join(str(cell or "").split()) for cell in row]
        if any(cleaned_row):
            cleaned_rows.append(cleaned_row)

    if not cleaned_rows:
        return []

    num_cols = max(len(row) for row in cleaned_rows)
    normalized_rows = [row + [""] * (num_cols - len(row)) for row in cleaned_rows]

    header = normalized_rows[0]
    separator = ["---"] * num_cols

    if len(normalized_rows) == 1:
        md_lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(separator) + " |"
        ]
        return ["\n".join(md_lines)]

    chunks = []
    for i in range(1, len(normalized_rows), max_rows):
        batch = normalized_rows[i : i + max_rows]
        md_lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(separator) + " |"
        ]
        for row in batch:
            md_lines.append("| " + " | ".join(row) + " |")
        chunks.append("\n".join(md_lines))

    return chunks


def extract_text_and_tables(pdf_path):
    """
    Extracts structured content page-by-page from a PDF file using PyMuPDF:
    1. Detects and serializes tables as intact Markdown tables to preserve row-column semantics.
    2. Extracts standard narrative text.
    Returns: List of tuples (page_num, content_text, content_type)
    """
    doc = fitz.open(pdf_path)
    extracted_items = []

    for i, page in enumerate(doc):
        page_num = i + 1

        # 1. Extract Structured Tables
        try:
            tables = page.find_tables()
        except Exception:
            tables = []

        if tables and hasattr(tables, "tables"):
            for tab in tables.tables:
                table_data = tab.extract()
                md_tables = table_to_markdown_chunks(table_data)
                for idx, md_table in enumerate(md_tables):
                    part = f" (Part {idx+1})" if len(md_tables) > 1 else ""
                    extracted_items.append((
                        page_num,
                        f"[TABLE - Page {page_num}{part}]\n{md_table}\n[/TABLE]",
                        "table"
                    ))

        # 2. Extract Narrative Text
        text = page.get_text().strip()
        if text:
            extracted_items.append((page_num, text, "text"))

    doc.close()
    return extracted_items


def extract_text_from_excel(excel_path: str) -> List[tuple]:
    """
    Extracts structured content from an Excel file (.xlsx / .xls) using openpyxl.
    Each sheet is treated as a "page". Tables are serialized as Markdown tables.
    Returns: List of tuples (sheet_num, content_text, content_type)
    """
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    extracted_items = []

    for sheet_num, sheet_name in enumerate(wb.sheetnames, start=1):
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            # Skip completely empty rows
            if any(cell is not None for cell in row):
                rows.append([str(cell) if cell is not None else "" for cell in row])

        if not rows:
            continue

        # Serialize all rows into Markdown table chunks
        md_tables = table_to_markdown_chunks(rows, max_rows=30)
        for idx, md_table in enumerate(md_tables):
            part = f" (Part {idx+1})" if len(md_tables) > 1 else ""
            extracted_items.append((
                sheet_num,
                f"[TABLE - Sheet {sheet_num}: {sheet_name}{part}]\n{md_table}\n[/TABLE]",
                "table"
            ))

    wb.close()
    return extracted_items


def extract_text_from_docx(docx_path: str) -> List[tuple]:
    """
    Extracts structured content from a DOCX file using python-docx.
    Paragraphs are grouped into pages (by section breaks / every 30 paragraphs).
    Tables are serialized as Markdown tables.
    Returns: List of tuples (page_num, content_text, content_type)
    """
    doc = DocxDocument(docx_path)
    extracted_items = []
    page_num = 1
    paragraph_buffer = []
    PARAGRAPHS_PER_PAGE = 30  # approximate grouping when no section breaks

    # Extract tables first (with their approximate position)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        md_tables = table_to_markdown_chunks(rows)
        for idx, md_table in enumerate(md_tables):
            part = f" (Part {idx+1})" if len(md_tables) > 1 else ""
            extracted_items.append((
                page_num,
                f"[TABLE - Page {page_num}{part}]\n{md_table}\n[/TABLE]",
                "table"
            ))

    # Extract paragraph text, grouping into virtual pages
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraph_buffer.append(text)

        # Flush buffer every PARAGRAPHS_PER_PAGE paragraphs
        if len(paragraph_buffer) >= PARAGRAPHS_PER_PAGE:
            extracted_items.append((page_num, "\n".join(paragraph_buffer), "text"))
            paragraph_buffer = []
            page_num += 1

    # Flush remaining paragraphs
    if paragraph_buffer:
        extracted_items.append((page_num, "\n".join(paragraph_buffer), "text"))

    return extracted_items


MIN_IMAGE_WIDTH = 150
MIN_IMAGE_HEIGHT = 150
MAX_IMAGE_DIM = 800


def extract_pdf_images(pdf_path: str) -> List[Dict]:
    """
    Extracts significant images and charts from a PDF file.
    Filters out small icons, bullet graphics, and tiny logos (width/height < 150px).
    Resizes images to max 800px for ultra-fast Vision AI processing.
    """
    doc = fitz.open(pdf_path)
    extracted_images = []

    for i, page in enumerate(doc):
        page_num = i + 1
        try:
            image_list = page.get_images(full=True)
        except Exception:
            image_list = []

        for img_idx, img_info in enumerate(image_list):
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                width = base_image["width"]
                height = base_image["height"]

                # Ignore tiny icons/logos
                if width < MIN_IMAGE_WIDTH or height < MIN_IMAGE_HEIGHT:
                    continue

                pil_img = Image.open(io.BytesIO(image_bytes))
                if pil_img.mode in ("RGBA", "P"):
                    pil_img = pil_img.convert("RGB")

                # Scale down for fast payload transmission
                if max(pil_img.size) > MAX_IMAGE_DIM:
                    pil_img.thumbnail((MAX_IMAGE_DIM, MAX_IMAGE_DIM), Image.Resampling.LANCZOS)

                buffer = io.BytesIO()
                pil_img.save(buffer, format="JPEG", quality=85)
                b64_str = base64.b64encode(buffer.getvalue()).decode("utf-8")

                extracted_images.append({
                    "page": page_num,
                    "index": img_idx + 1,
                    "width": width,
                    "height": height,
                    "base64": b64_str,
                })
            except Exception:
                continue

    doc.close()
    return extracted_images


def describe_chart_image(base64_img: str, provider: str = "ollama", model: str = "llama3.2-vision", api_key: Optional[str] = None) -> str:
    """
    Analyzes an extracted chart or figure using Vision AI to extract numeric values, axes, and findings.
    Supports Google Gemini, OpenAI GPT-4o, Anthropic Claude, and local Ollama Vision.
    """
    prompt = (
        "Analyze this chart or diagram from a document. "
        "Extract all key findings, numerical values, percentages, axes labels, category names, "
        "and trends in concise bullet points. Focus purely on factual data without assumptions."
    )

    try:
        if provider == "gemini" and api_key:
            gemini_model = model if "gemini" in model else "gemini-2.5-flash"
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{gemini_model}:generateContent?key={api_key}"
            payload = {
                "contents": [{
                    "parts": [
                        {"text": prompt},
                        {"inlineData": {"mimeType": "image/jpeg", "data": base64_img}}
                    ]
                }]
            }
            resp = requests.post(url, json=payload, timeout=30)
            resp.raise_for_status()
            data = resp.json()
            return data["candidates"][0]["content"]["parts"][0]["text"].strip()

        elif provider == "openai" and api_key:
            client = OpenAI(api_key=api_key)
            oai_model = model if "gpt-4" in model else "gpt-4o-mini"
            resp = client.chat.completions.create(
                model=oai_model,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_img}"}}
                    ]
                }],
                max_tokens=300,
            )
            return resp.choices[0].message.content.strip()

        elif provider == "anthropic" and api_key:
            claude_model = model if "claude-3" in model else "claude-3-5-haiku-20241022"
            payload = {
                "model": claude_model,
                "max_tokens": 300,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": base64_img}}
                    ]
                }]
            }
            resp = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": api_key,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                },
                json=payload,
                timeout=30,
            )
            resp.raise_for_status()
            return resp.json()["content"][0]["text"].strip()

        else:
            # Local Ollama Vision
            ollama_model = model if ("vision" in model or "llava" in model) else "llama3.2-vision"
            resp = requests.post(
                f"{OLLAMA_URL}/api/generate",
                json={
                    "model": ollama_model,
                    "prompt": prompt,
                    "images": [base64_img],
                    "stream": False,
                    "options": {"num_predict": 300}
                },
                timeout=60,
            )
            resp.raise_for_status()
            return resp.json().get("response", "").strip()

    except Exception:
        return "[Visual Figure/Chart: contains visual workflow and graphical distribution data]"


def _split_sentences(text):
    """
    Two-level text splitting:
    1. Splits on paragraph boundaries (double newlines) first.
    2. Splits within each paragraph on sentence endings (. ! ? ؟).
    This ensures clean sentence boundaries for both Arabic and Latin scripts.
    """
    paragraphs = re.split(r'\n\s*\n', text)
    sentences = []
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        parts = re.split(r'(?<=[.!?؟])\s+', para)
        sentences.extend(s for s in parts if s)

    return sentences


def chunk_text(text):
    """
    Sentence-aware chunking with dynamic overlap and hard-split fallback.
    Preserves complete sentences per chunk and carries trailing sentences into adjacent chunks
    without cutting mid-sentence.
    """
    if not text or not text.strip():
        return []

    sentences = _split_sentences(text)
    if not sentences:
        return []

    chunks = []
    current_sentences = []
    current_length = 0

    for sentence in sentences:
        sentence_length = len(sentence)

        # Fallback hard-split for single sentences exceeding CHUNK_SIZE
        if sentence_length > CHUNK_SIZE:
            if current_sentences:
                chunks.append(" ".join(current_sentences))
                current_sentences = []
                current_length = 0
            start = 0
            while start < sentence_length:
                chunks.append(sentence[start : start + CHUNK_SIZE])
                start += CHUNK_SIZE - CHUNK_OVERLAP
            continue

        # Check if adding the sentence exceeds CHUNK_SIZE
        if current_length + sentence_length + (1 if current_sentences else 0) > CHUNK_SIZE:
            if current_sentences:
                chunks.append(" ".join(current_sentences))

                # Sentence-aware overlap: carry last 1-2 sentences
                overlap_sentences = current_sentences[-2:] if len(current_sentences) >= 2 else current_sentences[-1:]

                # Ensure overlap + new sentence does not overflow CHUNK_SIZE
                while overlap_sentences:
                    projected = sum(len(s) for s in overlap_sentences) + sentence_length + len(overlap_sentences)
                    if projected <= CHUNK_SIZE:
                        break
                    overlap_sentences.pop(0)

                current_sentences = list(overlap_sentences)
                current_length = sum(len(s) for s in current_sentences) + max(len(current_sentences) - 1, 0)
            else:
                current_sentences = []
                current_length = 0

        current_sentences.append(sentence)
        current_length += sentence_length + (1 if len(current_sentences) > 1 else 0)

    # Append remaining sentences
    if current_sentences:
        last_chunk = " ".join(current_sentences)
        if not chunks or last_chunk != chunks[-1]:
            chunks.append(last_chunk)

    return chunks


OLLAMA_URL = "http://localhost:11434"
EMBEDDING_MODEL = "nomic-embed-text"


def embed_texts(texts):
    """
    Generates local embeddings via Ollama nomic-embed-text.
    Embeddings always run locally to ensure zero-cost, offline ingestion.
    Includes robust retry logic and error handling for large files.
    """
    import time
    embeddings = []
    for i, text in enumerate(texts):
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = requests.post(
                    f"{OLLAMA_URL}/api/embeddings",
                    json={"model": EMBEDDING_MODEL, "prompt": text},
                    timeout=120
                )
                response.raise_for_status()
                embeddings.append(response.json()["embedding"])
                break
            except Exception as e:
                if attempt == max_retries - 1:
                    print(f"Failed to embed chunk {i} after {max_retries} attempts: {e}")
                    # Fallback to zero vector to prevent crashing the entire ingestion
                    embeddings.append([0.0] * 768)
                else:
                    print(f"Embedding attempt {attempt+1} failed, retrying in {2**attempt}s...")
                    time.sleep(2 ** attempt)
    return embeddings


DEFAULT_SYSTEM_PROMPT = """You are a highly capable, precise, and professional AI document analysis assistant.

Operating Guidelines:
1. STRICT LANGUAGE MATCHING: You MUST ALWAYS respond in the EXACT SAME LANGUAGE as the user's question. If asked in Arabic, respond ONLY in fluent, professional, modern standard Arabic. If asked in English, respond in clear English. NEVER switch languages or append English apologies to Arabic responses.
2. FACTUAL GROUNDING: Rely strictly on the provided context. Do NOT invent, assume, or fabricate facts beyond what is documented in the context.
3. INSUFFICIENT INFORMATION: If the answer cannot be determined from the context, state clearly and politely in the query language that the information is not present in the provided documents.
4. STRUCTURE & TONE: Maintain a polite, objective, and well-structured tone. Use clear markdown formatting (bullet points, bold text, concise paragraphs) for readability.
5. MERMAID DIAGRAMS: CRITICAL: You MUST wrap all text inside nodes with double quotes to prevent syntax errors (e.g., A["Artificial Intelligence (AI)"] or B["الذكاء الاصطناعي"]). Do not use unquoted spaces, parentheses, or Arabic text inside node brackets."""


COMPARE_SYSTEM_PROMPT = """You are an expert AI policy and document analyst specialized in comparative analysis and gap assessment.

Operating Guidelines:
1. Provide a rigorous, well-structured, and factual gap analysis comparing the two provided documents.
2. Structure your response clearly:
   - Topics/Policies unique to Document A
   - Topics/Policies unique to Document B
   - Shared topics with key similarities and differences
   - Strategic and actionable recommendations
3. Strict Factual Fidelity: Do not hallucinate or assume external policies. Cite the source document for each observation.
4. Language Matching: Answer in the same language as the user query or primary document context.
5. Visual Diagram: Conclude with a clean, valid Mermaid diagram in a ```mermaid ... ``` code block (such as graph TD or pie) summarizing the key comparative relationships or distributions visually. CRITICAL: You MUST wrap all text inside Mermaid nodes with double quotes to prevent syntax errors (e.g., A["Artificial Intelligence (AI)"] or B["الذكاء الاصطناعي"]). Do not use unquoted spaces, parentheses, HTML tags, or Arabic text inside node brackets."""


def generate_with_openai(prompt: str, model: str, api_key: str, system_prompt: Optional[str] = None, history: Optional[list] = None, base_url: Optional[str] = None) -> str:
    provider_client = OpenAI(api_key=api_key, base_url=base_url)
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    if history:
        for turn in history:
            if isinstance(turn, dict) and "role" in turn and "content" in turn:
                messages.append({"role": turn["role"], "content": turn["content"]})
    messages.append({"role": "user", "content": prompt})

    response = provider_client.chat.completions.create(
        model=model,
        messages=messages,
        timeout=120,
    )
    return response.choices[0].message.content


def generate_with_anthropic(prompt: str, model: str, api_key: str, system_prompt: Optional[str] = None, history: Optional[list] = None) -> str:
    messages = []
    if history:
        for turn in history:
            if isinstance(turn, dict) and "role" in turn and "content" in turn:
                messages.append({"role": turn["role"], "content": turn["content"]})
    messages.append({"role": "user", "content": prompt})

    payload = {
        "model": model,
        "max_tokens": 4096,
        "messages": messages,
    }
    if system_prompt:
        payload["system"] = system_prompt

    response = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
        json=payload,
        timeout=120,
    )
    response.raise_for_status()
    return response.json()["content"][0]["text"]


def generate_with_gemini(prompt: str, model: str, api_key: str, system_prompt: Optional[str] = None, history: Optional[list] = None) -> str:
    contents = []
    if history:
        for turn in history:
            if isinstance(turn, dict) and "role" in turn and "content" in turn:
                role = "user" if turn["role"] == "user" else "model"
                contents.append({"role": role, "parts": [{"text": turn["content"]}]})
    contents.append({"role": "user", "parts": [{"text": prompt}]})

    payload = {
        "contents": contents,
        "generationConfig": {"maxOutputTokens": 4096},
    }
    if system_prompt:
        payload["systemInstruction"] = {"parts": [{"text": system_prompt}]}

    response = requests.post(
        f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}",
        json=payload,
        timeout=120,
    )
    # If model not found (404) or forbidden (403), try fallbacks
    if response.status_code in [404, 403]:
        for fallback in ["gemini-pro-latest", "gemini-3.1-pro-preview", "gemini-2.5-flash"]:
            if fallback == model:
                continue
            response = requests.post(
                f"https://generativelanguage.googleapis.com/v1beta/models/{fallback}:generateContent?key={api_key}",
                json=payload,
                timeout=120,
            )
            if response.status_code not in [404, 403]:
                break
    # Handle rate limiting with retry
    if response.status_code == 429:
        import time as _time
        for wait in [5, 15, 30]:
            print(f"Gemini rate limited, retrying in {wait}s...")
            _time.sleep(wait)
            response = requests.post(
                f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}",
                json=payload,
                timeout=120,
            )
            if response.status_code != 429:
                break
    response.raise_for_status()
    return response.json()["candidates"][0]["content"]["parts"][0]["text"]


def generate_with_ollama(prompt: str, model: str, system_prompt: Optional[str] = None, history: Optional[list] = None, max_tokens: int = 2048) -> str:
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    if history:
        for turn in history:
            if isinstance(turn, dict) and "role" in turn and "content" in turn:
                messages.append({"role": turn["role"], "content": turn["content"]})
    messages.append({"role": "user", "content": prompt})

    response = requests.post(
        f"{OLLAMA_URL}/api/chat",
        json={
            "model": model,
            "messages": messages,
            "stream": False,
            "options": {"num_predict": max_tokens, "temperature": 0.2, "top_p": 0.9},
        },
        timeout=6000,
    )
    response.raise_for_status()
    return response.json()["message"]["content"]


def generate_answer(prompt: str, provider: str, model: str, api_key: Optional[str], system_prompt: Optional[str] = None, history: Optional[list] = None, max_tokens: int = 2048) -> str:
    """Routes generation request to the selected provider with system prompt and conversation history support."""
    if provider == "ollama":
        return generate_with_ollama(prompt, model, system_prompt=system_prompt, history=history, max_tokens=max_tokens)
    if provider == "openai":
        return generate_with_openai(prompt, model, api_key, system_prompt=system_prompt, history=history)
    if provider == "deepseek":
        return generate_with_openai(prompt, model, api_key, system_prompt=system_prompt, history=history, base_url="https://api.deepseek.com")
    if provider == "moonshot":
        return generate_with_openai(prompt, model, api_key, system_prompt=system_prompt, history=history, base_url="https://api.moonshot.cn/v1")
    if provider == "anthropic":
        return generate_with_anthropic(prompt, model, api_key, system_prompt=system_prompt, history=history)
    if provider == "gemini":
        return generate_with_gemini(prompt, model, api_key, system_prompt=system_prompt, history=history)
    raise ValueError(f"Unknown provider: {provider}")


ALLOWED_EXTENSIONS = {".pdf", ".xlsx", ".xls", ".docx"}


@app.post("/ingest")
async def ingest(
    file: UploadFile = File(...),
    enable_vision: bool = Form(False),
    vision_provider: str = Form("ollama"),
    vision_model: str = Form("llama3.2-vision"),
    vision_api_key: Optional[str] = Form(None),
):
    """Uploads, extracts text, structured tables, and optionally visual charts into ChromaDB with concurrency protection.
    Supports PDF, Excel (.xlsx/.xls), and Word (.docx) files.
    """
    # Validate file extension
    _, file_ext = os.path.splitext(file.filename or "")
    file_ext = file_ext.lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{file_ext}'. Allowed types: PDF, Excel (.xlsx/.xls), Word (.docx)."
        )

    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    async with chroma_write_lock:
        collection = get_or_create_collection()

        # Route extraction based on file type
        if file_ext == ".pdf":
            items = extract_text_and_tables(tmp_path)
        elif file_ext in (".xlsx", ".xls"):
            items = extract_text_from_excel(tmp_path)
        elif file_ext == ".docx":
            items = extract_text_from_docx(tmp_path)
        else:
            items = []
        all_chunks = []
        all_metadata = []
        all_ids = []
        
        # Deterministic IDs
        safe_filename = "".join([c if c.isalnum() else "_" for c in file.filename])
        local_chunk_idx = 0

        for page_num, content_text, content_type in items:
            if content_type == "table":
                all_chunks.append(content_text)
                all_metadata.append({
                    "source": file.filename,
                    "page": page_num,
                    "content_type": "table",
                })
                all_ids.append(f"{safe_filename}_p{page_num}_{local_chunk_idx}")
                local_chunk_idx += 1
            else:
                for chunk in chunk_text(content_text):
                    if len(chunk.strip()) < 20:
                        continue
                    all_chunks.append(chunk)
                    all_metadata.append({
                        "source": file.filename,
                        "page": page_num,
                        "content_type": "text",
                    })
                    all_ids.append(f"{safe_filename}_p{page_num}_{local_chunk_idx}")
                    local_chunk_idx += 1

        # Process Visual Charts with Vision AI if requested (PDF only)
        figures_count = 0
        if enable_vision and file_ext == ".pdf":
            extracted_images = extract_pdf_images(tmp_path)
            for img in extracted_images:
                page_num = img["page"]
                fig_idx = img["index"]
                description = describe_chart_image(
                    img["base64"],
                    provider=vision_provider,
                    model=vision_model,
                    api_key=vision_api_key,
                )
                fig_chunk = f"[FIGURE - Page {page_num} - Chart #{fig_idx}]\n{description}\n[/FIGURE]"
                all_chunks.append(fig_chunk)
                all_metadata.append({
                    "source": file.filename,
                    "page": page_num,
                    "content_type": "figure",
                })
                all_ids.append(f"{safe_filename}_p{page_num}_{local_chunk_idx}")
                local_chunk_idx += 1
                figures_count += 1

        chunks_added = 0
        # Checkpoint Resumption & Batch Upsert Logic
        for i in range(0, len(all_chunks), 100):
            batch_texts = all_chunks[i : i + 100]
            batch_meta = all_metadata[i : i + 100]
            batch_ids = all_ids[i : i + 100]
            
            # Check which IDs already exist in ChromaDB
            existing = collection.get(ids=batch_ids)
            existing_ids = set(existing.get("ids", []))
            
            # Filter to only the new un-embedded chunks
            new_texts = []
            new_meta = []
            new_ids = []
            
            for t, m, i_d in zip(batch_texts, batch_meta, batch_ids):
                if i_d not in existing_ids:
                    new_texts.append(t)
                    new_meta.append(m)
                    new_ids.append(i_d)
                    
            if new_texts:
                # Embed and save only the new chunks
                new_embeddings = embed_texts(new_texts)
                collection.upsert(
                    documents=new_texts,
                    embeddings=new_embeddings,
                    metadatas=new_meta,
                    ids=new_ids,
                )
                chunks_added += len(new_texts)
                print(f"Ingested {len(new_texts)} new chunks (skipped {len(batch_texts) - len(new_texts)} existing).")
            else:
                print(f"Skipped {len(batch_texts)} existing chunks. Resuming...")

        total_chunks = collection.count()

    if os.path.exists(tmp_path):
        os.unlink(tmp_path)

    return {
        "message": f"Ingested {file.filename} successfully",
        "chunks_added": chunks_added,
        "figures_analyzed": figures_count,
        "total_chunks": total_chunks,
    }


# === Query Routing & Dynamic Conversational Onboarding ===
GREETING_PHRASES = [
    'مرحبا', 'اهل', 'اهلا', 'اهلين', 'هلا', 'سلام', 'السلام عليكم', 'وعليكم السلام',
    'تحياتي', 'صباح الخير', 'صباح النور', 'صباح الورد', 'مساء الخير', 'مساء النور', 'مساء الورد',
    'كيف حالك', 'كيفك', 'شلونك', 'اخبارك', 'شخبارك', 'كيف الصحه',
    'شكرا', 'مشكور', 'تسلم', 'يعطيك العافيه', 'جزاك الله',
    'من انت', 'ماذا تفعل', 'عرف عن نفسك',
    'hello', 'hi', 'hey', 'greetings', 'good morning', 'good afternoon', 'good evening',
    'how are you', 'how do you do', 'whats up', 'what is up', 'thank you', 'thanks',
    'who are you', 'what can you do'
]
MAX_GREETING_WORDS = 7


def normalize_text_for_routing(text: str) -> str:
    """Normalizes characters, diacritics, and punctuation for reliable intent matching."""
    t = text.lower().strip()
    t = re.sub(r'[\u064B-\u065F\u0670]', '', t)
    t = re.sub(r'[أإآ]', 'ا', t)
    t = re.sub(r'ة\b', 'ه', t)
    t = re.sub(r'ى\b', 'ي', t)
    t = re.sub(r'[ـ!؟?.,،؛:~\-_]+', ' ', t)
    return ' '.join(t.split())


def is_arabic(text: str) -> bool:
    """Checks whether the text contains Arabic characters."""
    return bool(re.search(r'[\u0600-\u06FF]', text))


def is_greeting(text: str) -> bool:
    """Detects whether the input message is a greeting or brief small-talk."""
    cleaned = normalize_text_for_routing(text)
    if not cleaned:
        return False
    words = cleaned.split()
    if len(words) > MAX_GREETING_WORDS:
        return False
    for phrase in GREETING_PHRASES:
        norm_p = normalize_text_for_routing(phrase)
        if cleaned == norm_p or cleaned.startswith(norm_p + ' ') or cleaned.endswith(' ' + norm_p) or (' ' + norm_p + ' ') in (' ' + cleaned + ' '):
            return True
    return False


# Stage 2: Relevance threshold — filters out off-topic queries while supporting cross-lingual and summary matches
RELEVANCE_THRESHOLD = 12.0

NO_RELEVANCE_RESPONSE_AR = "لم أجد معلومات ذات صلة كافية بمستنداتك المرفوعة لهذا السؤال."
NO_RELEVANCE_RESPONSE_EN = "I could not find sufficient relevant information in your uploaded documents for this question."


def get_no_relevance_response(text: str) -> str:
    """Returns a graceful refusal response in the matching query language."""
    return NO_RELEVANCE_RESPONSE_AR if is_arabic(text) else NO_RELEVANCE_RESPONSE_EN


def is_summary_request(text: str) -> bool:
    """Detects whether the user is asking for a comprehensive summary, overview, or main objectives/findings."""
    norm = normalize_text_for_routing(text)
    summary_terms = [
        'summary', 'summarize', 'overview', 'key summary', 'brief overview', 'outline',
        'objectives', 'main findings', 'overall', 'goals',
        'لخص', 'تلخيص', 'ملخص', 'نبذه', 'موجز', 'نظره عامه', 'نظرة عامة',
        'اهداف', 'الاهداف', 'ابرز الاهداف', 'نتائج', 'النتائج', 'ابرز النتائج', 'موضوع المستند'
    ]
    return any(term in norm for term in summary_terms)


def detect_target_documents(question: str, doc_names: list):
    """
    Detects whether the question targets a specific document by name/index,
    or requests a multi-document comparison.
    Returns: ('target', doc_name) | ('compare', doc_names) | ('global', None)
    """
    q_norm = normalize_text_for_routing(question)
    q_lower = question.lower()

    # Multi-document comparison intent
    compare_terms = [
        'compare', 'comparison', 'versus', 'vs', 'both documents',
        'قارن', 'مقارنه', 'مقارنة', 'كلا الملفين', 'الملفين', 'بين الملفين', 'بين المستندين'
    ]
    if any(term in q_norm for term in compare_terms) and len(doc_names) >= 2:
        return 'compare', doc_names

    # Check for direct doc filename or substring mentions
    for name in doc_names:
        name_clean = name.lower()
        base_name = name_clean.rsplit('.', 1)[0]
        if name_clean in q_lower or (len(base_name) >= 4 and base_name in q_lower):
            return 'target', name
        if 'jegh' in name_clean and 'jegh' in q_lower:
            return 'target', name
        if 'saudi_healthcare' in name_clean and any(t in q_lower for t in ['saudi', 'healthcare', 'الصح']):
            return 'target', name

    # Check for ordinal mentions
    if len(doc_names) >= 1 and any(p in q_norm for p in ['الملف الاول', 'المستند الاول', 'الملف 1', 'first document', 'doc 1', 'document 1']):
        return 'target', doc_names[0]

    if len(doc_names) >= 2 and any(p in q_norm for p in ['الملف الثاني', 'المستند الثاني', 'الملف 2', 'second document', 'doc 2', 'document 2']):
        return 'target', doc_names[1]

    return 'global', None


def generate_dynamic_greeting(collection, user_message: str, provider: str, model: str, api_key: Optional[str]) -> str:
    """
    Generates a conversational onboarding greeting.
    Uses the LLM to generate highly relevant, context-aware starter questions based on actual document chunks.
    """
    is_ar = is_arabic(user_message)

    if collection.count() == 0:
        if is_ar:
            return (
                "وعليكم السلام ورحمة الله وبركاته! أهلاً بك في نظام RAG الذكي. 📚\n\n"
                "لم تقم برفع أي مستندات بعد في قاعدة المعرفة. يمكنك رفع ملف PDF من القائمة الجانبية (Sidebar) "
                "وسأقوم بقراءته وتحليله ومناقشة محتواه معك خطوة بخطوة."
            )
        else:
            return (
                "Hello and welcome to your AI Document Assistant! 📚\n\n"
                "No documents have been uploaded to the knowledge base yet. Please upload a document using the sidebar "
                "so we can analyze and explore it together."
            )

    # Retrieve a diverse sample of documents to generate smart questions
    meta_items = collection.get(include=["metadatas"])
    doc_names = list(dict.fromkeys(m.get("source", "Document") for m in meta_items.get("metadatas", [])))

    sample_texts = []
    for source in doc_names[:3]:
        res = collection.get(where={"source": source}, include=["documents"], limit=1)
        if res and res.get("documents"):
            sample_texts.append(f"[Document: {source}]\n{res['documents'][0]}")

    smart_questions = ""
    try:
        context_sample = "\n\n".join(sample_texts)
        
        if is_ar:
            prompt = f"بناءً على هذه المقتطفات من المستندات:\n{context_sample}\n\nاقترح 3 أسئلة مفيدة ومحددة جداً يمكن للمستخدم طرحها لاستكشاف هذا المحتوى. اكتب الأسئلة فقط في قائمة مرقمة (1. ، 2. ، 3.) بدون أي مقدمات أو خاتمة."
            sys_prompt = "أنت مساعد ذكي. أخرج قائمة مرقمة بالأسئلة فقط."
        else:
            prompt = f"Based on these document excerpts:\n{context_sample}\n\nSuggest 3 specific and insightful questions the user could ask to explore this content. Write only the questions as a numbered list (1., 2., 3.) without any intro or outro."
            sys_prompt = "You are a helpful assistant. Output only the requested numbered list."
            
        smart_questions = generate_answer(
            prompt, provider, model, api_key,
            system_prompt=sys_prompt,
            max_tokens=150
        )
    except Exception as e:
        # Fallback to hardcoded generic questions if LLM fails
        if is_ar:
            q1 = f"ما هي الأهداف والنتائج الرئيسية المذكورة في {doc_names[0]}؟" if doc_names else "ما هي الأهداف الرئيسية؟"
            q2 = f"ما هي أهم الحلول والبيانات المطروحة في {doc_names[1]}؟" if len(doc_names) > 1 else "ما هي أهم النتائج والبيانات المذكورة في المستند؟"
            q3 = "قارن بين المستندات المرفوعة واستنتج رسماً بيانياً توضيحياً للفروقات." if len(doc_names) > 1 else "لخص لي أهم محاور وجداول هذا المستند."
            smart_questions = f"1. {q1}\n2. {q2}\n3. {q3}"
        else:
            q1 = f"What are the main objectives and findings in {doc_names[0]}?" if doc_names else "What are the main objectives?"
            q2 = f"How does {doc_names[1]} address its core topic?" if len(doc_names) > 1 else "What are the practical applications and recommendations mentioned?"
            q3 = "Compare the uploaded documents and generate a visual chart." if len(doc_names) > 1 else "Summarize the key data and tables in the document."
            smart_questions = f"1. {q1}\n2. {q2}\n3. {q3}"

    if is_ar:
        docs_str = "\n".join([f"- 📄 **{name}**" for name in doc_names])
        return (
            f"وعليكم السلام ورحمة الله وبركاته! أهلاً وسهلاً بك. 📚✨\n\n"
            f"**المستندات المتاحة حالياً في قاعدة المعرفة:**\n{docs_str}\n\n"
            f"**أسئلة مقترحة لبدء النقاش (بناءً على محتوى ملفاتك):**\n"
            f"{smart_questions.strip()}\n\n"
            f"تفضل بطرح أي سؤال للبدء في استكشاف وتحليل المحتوى!"
        )
    else:
        docs_str = "\n".join([f"- 📄 **{name}**" for name in doc_names])
        return (
            f"Hello and welcome! 📚✨\n\n"
            f"**Documents currently in your knowledge base:**\n{docs_str}\n\n"
            f"**Suggested starter questions (based on your documents):**\n"
            f"{smart_questions.strip()}\n\n"
            f"Feel free to ask any question to begin exploring!"
        )


class AskRequest(BaseModel):
    question: str
    n_results: int = 5
    provider: str = "ollama"               # "ollama" | "openai" | "anthropic" | "gemini"
    model: str = "llama3.1"                # Target model identifier
    api_key: Optional[str] = None          # Optional API key for cloud providers
    history: Optional[list] = None         # Prior conversation turns for multi-turn context
    session_id: Optional[str] = None       # Active conversation session ID
    quoted_message: Optional[str] = None   # Text of specific previous message being replied to


@app.post("/ask")
async def ask(request: AskRequest):
    """
    Executes the RAG query pipeline with:
    1. Stage 1: Greeting & Conversational Onboarding fast-path.
    2. Embedding and ChromaDB vector retrieval (with multi-doc summary aggregation for overview queries).
    3. Stage 2: Relevance threshold filter to prevent hallucinations.
    4. Multi-provider LLM generation with multi-turn conversational memory, reply quoting, and session persistence.
    """
    collection = get_or_create_collection()

    # Stage 1: Fast-path for greetings & dynamic discussion starters
    if is_greeting(request.question):
        dynamic_greeting = generate_dynamic_greeting(
            collection=collection,
            user_message=request.question,
            provider=request.provider,
            model=request.model,
            api_key=request.api_key,
        )

        if request.session_id:
            async with sessions_lock:
                try:
                    session = get_session_by_id(request.session_id)
                    title = session["title"] if (session and session.get("title")) else (request.question[:35] + "..." if len(request.question) > 35 else request.question)
                    existing_msgs = session["messages"] if session else []
                    existing_msgs.append({
                        "role": "user",
                        "content": request.question,
                        "quoted_message": request.quoted_message,
                    })
                    existing_msgs.append({
                        "role": "assistant",
                        "content": dynamic_greeting,
                        "sources": [],
                    })
                    save_or_update_session(request.session_id, title, existing_msgs)
                except Exception:
                    pass

        return {
            "question": request.question,
            "answer": dynamic_greeting,
            "sources": [],
        }

    if collection.count() == 0:
        return {"error": "No documents ingested yet"}

    # Filter and preserve a clean sliding window of the last 6 messages (last 3 full turns)
    valid_history = []
    if request.history:
        for turn in request.history[-6:]:
            if isinstance(turn, dict) and "role" in turn and "content" in turn:
                valid_history.append({"role": turn["role"], "content": turn["content"]})

    all_items = collection.get(include=["metadatas"])
    doc_names = list(dict.fromkeys(m.get("source", "Document") for m in all_items["metadatas"]))

    # Handle global summarization requests across all uploaded documents
    if is_summary_request(request.question):
        all_doc_items = collection.get(include=["documents", "metadatas"])
        doc_samples = {}
        for doc_text, meta in zip(all_doc_items["documents"], all_doc_items["metadatas"]):
            source = meta.get("source", "Document")
            page = meta.get("page", 1)
            if source not in doc_samples:
                doc_samples[source] = []
            if len(doc_samples[source]) < 2:
                doc_samples[source].append((page, doc_text))

        context_parts = []
        sources = []
        for doc_name, samples in doc_samples.items():
            for page, text in samples:
                context_parts.append(f"=== Document: {doc_name} (Page {page}) ===\n{text}")
                sources.append({
                    "source": doc_name,
                    "page": page,
                    "score": 100.0,
                    "content_type": "text",
                })

        context = "\n\n".join(context_parts)
    else:
        question_embedding = embed_texts([request.question])[0]
        intent, target = detect_target_documents(request.question, doc_names)

        if intent == 'target' and target:
            # Query targeted document specifically
            results = collection.query(
                query_embeddings=[question_embedding],
                n_results=min(request.n_results, collection.count()),
                where={"source": target},
                include=["documents", "metadatas", "distances"]
            )
            chunks = results["documents"][0] if results["documents"] else []
            metadatas = results["metadatas"][0] if results["metadatas"] else []
            distances = results["distances"][0] if results["distances"] else []

        elif intent == 'compare' and len(doc_names) >= 2:
            # Balanced multi-document retrieval across each document
            chunks, metadatas, distances = [], [], []
            per_doc_k = max(2, request.n_results // len(doc_names) + 1)
            for d_name in doc_names:
                doc_res = collection.query(
                    query_embeddings=[question_embedding],
                    n_results=per_doc_k,
                    where={"source": d_name},
                    include=["documents", "metadatas", "distances"]
                )
                if doc_res["documents"] and doc_res["documents"][0]:
                    chunks.extend(doc_res["documents"][0])
                    metadatas.extend(doc_res["metadatas"][0])
                    distances.extend(doc_res["distances"][0])

        else:
            # Global retrieval across all documents
            results = collection.query(
                query_embeddings=[question_embedding],
                n_results=request.n_results,
                include=["documents", "metadatas", "distances"]
            )
            chunks = results["documents"][0] if results["documents"] else []
            metadatas = results["metadatas"][0] if results["metadatas"] else []
            distances = results["distances"][0] if results["distances"] else []

        context_parts = []
        sources = []

        for chunk, meta, dist in zip(chunks, metadatas, distances):
            context_parts.append(chunk)
            similarity_pct = round(100 * (2.718281828 ** (-dist / 220)), 1)
            sources.append({
                "source": meta["source"],
                "page": meta["page"],
                "score": similarity_pct,
                "content_type": meta.get("content_type", "text"),
            })

        # Stage 2: Low-relevance filter — rejects off-topic queries without wasting LLM calls
        best_score = max(s["score"] for s in sources) if sources else 0
        if best_score < RELEVANCE_THRESHOLD and not valid_history and not request.quoted_message:
            return {
                "question": request.question,
                "answer": get_no_relevance_response(request.question),
                "sources": sources,
            }

        context = "\n\n".join(context_parts)

    quoted_context = ""
    if request.quoted_message:
        quoted_context = f"\n=== Replying Directly to Specific Previous Message ===\n\"{request.quoted_message}\"\n===\n"

    if is_arabic(request.question):
        lang_directive = (
            "- LANGUAGE: Write your ENTIRE response in fluent, professional Modern Standard Arabic (اللغة العربية الفصحى فقط).\n"
            "- CROSS-LINGUAL EXTRACTION: The document excerpts above are in English. Thoroughly extract the facts, findings, numbers, and details from them and explain them clearly in Arabic.\n"
            "- STRUCTURE: Provide a direct, well-organized response with clear bullet points and bold headings based strictly on the context."
        )
    else:
        lang_directive = (
            "- LANGUAGE: Write your response in clear, professional English.\n"
            "- STRUCTURE: Provide a direct, well-organized response with clear bullet points strictly based on the context above."
        )

    prompt = f"""Context from uploaded documents:
===
{context}
===
{quoted_context}
User Question: {request.question}

Instructions:
{lang_directive}
- If replying directly to a specific previous message quoted above, prioritize addressing that specific point."""

    if request.provider != "ollama" and not request.api_key:
        return {"error": f"Provider '{request.provider}' requires an API key. Please enter it in the sidebar."}

    async with generation_semaphore:
        try:
            answer_text = generate_answer(
                prompt,
                request.provider,
                request.model,
                request.api_key,
                system_prompt=DEFAULT_SYSTEM_PROMPT,
                history=valid_history if valid_history else None,
            )
        except Exception as e:
            return {"error": f"Generation failed via {request.provider}: {e}"}

    # Auto-persist conversation turn to SQLite session atomically
    if request.session_id:
        async with sessions_lock:
            try:
                session = get_session_by_id(request.session_id)
                title = session["title"] if (session and session.get("title")) else (request.question[:35] + "..." if len(request.question) > 35 else request.question)
                existing_msgs = session["messages"] if session else []
                existing_msgs.append({
                    "role": "user",
                    "content": request.question,
                    "quoted_message": request.quoted_message,
                })
                existing_msgs.append({
                    "role": "assistant",
                    "content": answer_text,
                    "sources": sources,
                })
                save_or_update_session(request.session_id, title, existing_msgs)
            except Exception:
                pass

    return {
        "question": request.question,
        "answer": answer_text,
        "sources": sources,
    }


# === Chat Sessions REST Endpoints with Concurrency Lock ===
@app.get("/sessions")
async def list_sessions_endpoint():
    """Returns a list of all saved chat sessions sorted by newest activity."""
    async with sessions_lock:
        return {"sessions": get_all_sessions()}


@app.get("/sessions/{session_id}")
async def get_session_endpoint(session_id: str):
    """Retrieves a specific chat session with its full message history."""
    async with sessions_lock:
        session = get_session_by_id(session_id)
        if not session:
            return {"error": f"Session '{session_id}' not found."}
        return {"session": session}


class SaveSessionRequest(BaseModel):
    id: str
    title: Optional[str] = None
    messages: List[Dict]


@app.post("/sessions")
async def save_session_endpoint(req: SaveSessionRequest):
    """Creates or updates a persistent conversation session."""
    async with sessions_lock:
        title = req.title or (req.messages[0]["content"][:35] if req.messages else "New Chat")
        save_or_update_session(req.id, title, req.messages)
        return {"message": "Session saved successfully", "id": req.id, "title": title}


@app.delete("/sessions/{session_id}")
async def delete_session_endpoint(session_id: str):
    """Deletes a conversation session from the SQLite database."""
    async with sessions_lock:
        deleted = delete_session_by_id(session_id)
        if not deleted:
            return {"error": f"Session '{session_id}' not found."}
        return {"message": f"Session '{session_id}' deleted successfully."}


@app.get("/documents")
def list_documents():
    """Returns a list of all ingested document names and their chunk counts."""
    collection = get_or_create_collection()
    if collection.count() == 0:
        return {"documents": []}
    all_items = collection.get(include=["metadatas"])
    counts = {}
    for meta in all_items["metadatas"]:
        source = meta["source"]
        counts[source] = counts.get(source, 0) + 1
    return {"documents": [{"name": name, "chunks": n} for name, n in counts.items()]}


@app.delete("/documents/{doc_name:path}")
async def delete_document(doc_name: str):
    """
    Deletes a specific document and all its associated chunks from the knowledge base atomically.
    """
    async with chroma_write_lock:
        collection = get_or_create_collection()
        if collection.count() == 0:
            return {"error": "Knowledge base is empty, nothing to delete."}

        # Retrieve all chunk IDs that belong to this document
        results = collection.get(
            where={"source": doc_name},
            include=["metadatas"]
        )

        matching_ids = results.get("ids", [])
        if not matching_ids:
            return {"error": f"Document '{doc_name}' was not found in the knowledge base."}

        chunks_removed = len(matching_ids)

        # Delete all matching chunks in a single batch
        collection.delete(ids=matching_ids)

        return {
            "message": f"Successfully deleted '{doc_name}' and all its chunks.",
            "document": doc_name,
            "chunks_removed": chunks_removed,
            "total_remaining": collection.count(),
        }


class CompareRequest(BaseModel):
    doc_a: str
    doc_b: str
    provider: str = "ollama"
    model: str = "llama3.1"
    api_key: Optional[str] = None


MAX_CHUNKS_PER_DOC_FOR_COMPARE = 10


@app.post("/compare")
async def compare_policies(request: CompareRequest):
    """
    Policy Gap Analysis: Compares two ingested documents and identifies
    unique topics, shared differences, and actionable recommendations.
    """
    collection = get_or_create_collection()

    def get_doc_text(filename):
        result = collection.get(
            where={"source": filename},
            include=["documents", "metadatas"]
        )
        if not result["documents"]:
            return ""
        pairs = sorted(
            zip(result["documents"], result["metadatas"]),
            key=lambda p: p[1]["page"]
        )
        pairs = pairs[:MAX_CHUNKS_PER_DOC_FOR_COMPARE]
        return "\n\n".join(chunk for chunk, _ in pairs)

    text_a = get_doc_text(request.doc_a)
    text_b = get_doc_text(request.doc_b)

    if not text_a or not text_b:
        return {"error": "One of the documents was not found in the database."}

    dimensions_list = "\n".join(f"- {d}" for d in ITU_AI_READINESS_DIMENSIONS)

    prompt = f"""You are an AI policy analyst evaluating documents against the ITU AI Readiness 2.0 framework's 13 official dimensions:
{dimensions_list}

Compare the two policy/strategy documents below and produce a structured gap analysis. Structure your analysis using the exact dimension names above as section headers (bold). Under each dimension header, write 1-2 sentences of specific findings grounded in the actual document text — do not write generic filler like 'not mentioned' for every dimension; only include a dimension if the documents contain relevant content for it, and skip dimensions with no evidence in either document rather than listing them as empty. Also note if the gap relates to ITU-T Y.3172 pipeline nodes (SRC, Collector, Pre-producer, Model, Policy, Distributor, SINK) where relevant — this is most relevant to the 'AI & Policies' and 'Digital Infrastructure' dimensions.

Focus on: (1) topics covered in one document but missing in the other, (2) differences in depth or specificity on shared topics, (3) concrete recommendations to close the identified gaps. Cite which document each point comes from. Do not fabricate information not present in either document.

=== Document A: {request.doc_a} ===
{text_a}

=== Document B: {request.doc_b} ===
{text_b}

Provide your analysis in exactly this format, using STRICT Markdown pipe-tables:

## 1. Topics only in Document A

| Dimension | Evidence |
|---|---|
| Strategy Alignment | "Quote from document..." (Document A) |
| Digital Infrastructure | "Quote from document..." (Document A) |

## 2. Topics only in Document B

| Dimension | Evidence |
|---|---|
| Digital Infrastructure | "Quote from document..." (Document B) |

## 3. Shared topics with notable differences

| Dimension | Document A | Document B | Differences |
|---|---|---|---|
| Strategy Alignment | AI is tied to Vision 2030... | Health system vision without AI... | A integrates AI; B does not. |

## 4. Recommendations to close the gaps

| Dimension | Recommendation |
|---|---|
| AI & Policies | Draft a national AI governance framework... |

CRITICAL RULES:
- You MUST use the pipe character '|' for EVERY table row. Every single data row must start and end with '|'.
- Include the separator line '|---|---|' after each header row.
- Do NOT output plain-text columnar layouts. ONLY use pipe-delimited Markdown tables.
- If quoting Arabic text from documents, provide an English translation in parentheses after the Arabic quote.
- Keep each table cell on a SINGLE line (no line breaks inside cells).
"""

    if request.provider != "ollama" and not request.api_key:
        return {"error": f"Provider '{request.provider}' requires an API key. Please enter it in the sidebar."}

    async with generation_semaphore:
        try:
            analysis = generate_answer(
                prompt,
                request.provider,
                request.model,
                request.api_key,
                system_prompt=COMPARE_SYSTEM_PROMPT
            )
        except Exception as e:
            return {"error": f"Analysis failed via {request.provider}: {e}"}

    return {"doc_a": request.doc_a, "doc_b": request.doc_b, "analysis": analysis}


last_known_local_models = []

@app.get("/models")
def available_models():
    """
    Discovers installed local Ollama models automatically,
    and returns supported cloud providers and standard models.
    """
    global last_known_local_models
    local_models = []
    try:
        resp = requests.get(f"{OLLAMA_URL}/api/tags", timeout=5)
        resp.raise_for_status()
        local_models = [m["name"] for m in resp.json().get("models", [])]
        last_known_local_models = local_models
    except Exception:
        local_models = last_known_local_models

    cloud_providers = {
        "openai": ["gpt-4o-mini", "gpt-4o", "gpt-4-turbo"],
        "anthropic": ["claude-3-5-sonnet-20241022", "claude-3-5-haiku-20241022", "claude-3-opus-20240229"],
        "gemini": ["gemini-pro-latest", "gemini-3.1-pro-preview", "gemini-2.5-flash", "gemini-2.5-pro"],
        "deepseek": ["deepseek-chat", "deepseek-reasoner"],
        "moonshot": ["moonshot-v1-8k", "moonshot-v1-32k", "moonshot-v1-128k"],
    }

    return {
        "local_models": local_models,
        "local_available": len(local_models) > 0,
        "cloud_providers": cloud_providers,
    }


class VerifyKeyRequest(BaseModel):
    provider: str
    api_key: str
    model: Optional[str] = None


@app.post("/verify-key")
async def verify_key_endpoint(request: VerifyKeyRequest):
    """
    Verifies that the provided API key is valid and connected to the target cloud provider.
    Returns connection status and token context window limits.
    """
    provider = request.provider.lower()
    api_key = request.api_key.strip()
    model = request.model or ""

    if not api_key:
        return {"status": "error", "message": "API key cannot be empty"}

    token_limits = {
        "openai": {"context_tokens": "128,000 Tokens", "tier": "GPT-4o Enterprise Reasoning"},
        "anthropic": {"context_tokens": "200,000 Tokens", "tier": "Claude 3.5 Sonnet Precision"},
        "gemini": {"context_tokens": "1,000,000 Tokens", "tier": "Gemini 1.5 Pro Extreme Context"},
        "deepseek": {"context_tokens": "64,000 Tokens", "tier": "DeepSeek V3 / Reasoner"},
        "moonshot": {"context_tokens": "128,000 Tokens", "tier": "Moonshot Kimi Context"},
    }

    try:
        if provider in ["openai", "deepseek", "moonshot"]:
            base_url = None
            if provider == "deepseek":
                base_url = "https://api.deepseek.com"
            elif provider == "moonshot":
                base_url = "https://api.moonshot.cn/v1"
            
            client = OpenAI(api_key=api_key, base_url=base_url)
            client.models.list()
            
            limit_info = token_limits.get(provider, {"context_tokens": "128,000 Tokens", "tier": "Cloud Model"})
            provider_name = provider.capitalize() if provider != "openai" else "OpenAI"
            
            return {
                "status": "connected",
                "provider": provider_name,
                "message": f"Successfully connected to {provider_name}! ({limit_info['tier']})",
                "context_tokens": limit_info["context_tokens"],
            }
        elif provider == "anthropic":
            test_resp = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": api_key,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                },
                json={
                    "model": model if model else "claude-3-5-sonnet-20241022",
                    "max_tokens": 5,
                    "messages": [{"role": "user", "content": "ping"}],
                },
                timeout=10,
            )
            if test_resp.status_code in [200, 400]:
                limit_info = token_limits.get("anthropic", {"context_tokens": "200,000 Tokens", "tier": "Anthropic Cloud"})
                return {
                    "status": "connected",
                    "provider": "Anthropic",
                    "message": f"Successfully connected to Anthropic! ({limit_info['tier']})",
                    "context_tokens": limit_info["context_tokens"],
                }
            else:
                return {"status": "error", "message": f"Anthropic key validation failed (Status {test_resp.status_code})"}
        elif provider == "gemini":
            # Use models.list endpoint for reliable key validation (no 404 risk from model names)
            test_resp = requests.get(
                f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}",
                timeout=10,
            )
            if test_resp.status_code == 200:
                # Parse available model names for display
                available_models = [m.get("name", "").replace("models/", "") for m in test_resp.json().get("models", [])]
                gemini_models = [m for m in available_models if "gemini" in m][:5]
                limit_info = token_limits.get("gemini", {"context_tokens": "1,000,000 Tokens", "tier": "Google Gemini Cloud"})
                return {
                    "status": "connected",
                    "provider": "Google Gemini",
                    "message": f"Successfully connected to Google Gemini! ({limit_info['tier']})",
                    "context_tokens": limit_info["context_tokens"],
                    "available_models": gemini_models,
                }
            else:
                return {"status": "error", "message": f"Google Gemini key validation failed (Status {test_resp.status_code})"}
        else:
            return {"status": "error", "message": f"Unsupported cloud provider: '{provider}'"}
    except Exception as e:
        return {"status": "error", "message": f"Cloud connection failed: {str(e)}"}


@app.get("/")
def root():
    """Returns API health status and vector database item count."""
    collection = get_or_create_collection()
    return {
        "status": "running",
        "total_chunks": collection.count(),
    }
